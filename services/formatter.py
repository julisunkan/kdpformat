"""
KDP Formatter Service
Main formatting logic for KDP-standard manuscript preparation.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from services.frontmatter import insert_front_matter
from services.dpi_checker import check_image_dpi


TRIM_SIZES = {
    '6x9': {'width': Inches(6), 'height': Inches(9)},
    '5x8': {'width': Inches(5), 'height': Inches(8)},
    '8.5x11': {'width': Inches(8.5), 'height': Inches(11)}
}

MARGINS = {
    'top': Inches(1),
    'bottom': Inches(1),
    'inside': Inches(0.85),
    'outside': Inches(0.6)
}


def format_manuscript(input_path, output_path, trim_size='6x9', print_mode=False, 
                      title="Untitled", author="Author Name", line_spacing=1.15):
    """
    Format a DOCX manuscript according to KDP standards.
    
    Args:
        input_path: Path to input DOCX file
        output_path: Path for output formatted DOCX
        trim_size: One of '6x9', '5x8', '8.5x11'
        print_mode: Enable mirrored margins for print
        title: Book title for front matter
        author: Author name for front matter
        line_spacing: Line spacing multiplier (1.15-1.5)
    
    Returns:
        Dictionary with status and any warnings
    """
    try:
        doc = Document(input_path)
        
        dpi_warnings = check_image_dpi(input_path)
        
        _setup_page_dimensions(doc, trim_size, print_mode)
        
        _setup_styles(doc, line_spacing)
        
        _cleanup_text(doc)
        
        _format_chapters(doc)
        
        _apply_body_formatting(doc, line_spacing)
        
        insert_front_matter(doc, title, author)
        
        _insert_toc(doc)
        
        doc.save(output_path)
        
        return {
            'success': True,
            'output_path': output_path,
            'dpi_warnings': dpi_warnings,
            'message': 'Manuscript formatted successfully'
        }
    
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'dpi_warnings': []
        }


def _setup_page_dimensions(doc, trim_size, print_mode):
    """Configure page size and margins."""
    size = TRIM_SIZES.get(trim_size, TRIM_SIZES['6x9'])
    
    for section in doc.sections:
        section.page_width = size['width']
        section.page_height = size['height']
        
        section.top_margin = MARGINS['top']
        section.bottom_margin = MARGINS['bottom']
        
        if print_mode:
            section.left_margin = MARGINS['inside']
            section.right_margin = MARGINS['outside']
            section.gutter = Inches(0)
            
            sectPr = section._sectPr
            mirrorMargins = OxmlElement('w:mirrorMargins')
            sectPr.append(mirrorMargins)
        else:
            section.left_margin = MARGINS['inside']
            section.right_margin = MARGINS['outside']


def _setup_styles(doc, line_spacing):
    """Set up document styles for body text and headings."""
    try:
        normal_style = doc.styles['Normal']
    except KeyError:
        normal_style = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    
    normal_style.font.name = 'Georgia'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = line_spacing
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.first_line_indent = Inches(0.25)
    
    try:
        heading1_style = doc.styles['Heading 1']
    except KeyError:
        heading1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    heading1_style.font.name = 'Georgia'
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_style.paragraph_format.space_before = Pt(24)
    heading1_style.paragraph_format.space_after = Pt(12)
    heading1_style.paragraph_format.first_line_indent = Inches(0)
    heading1_style.paragraph_format.page_break_before = True


def _cleanup_text(doc):
    """Remove extra spaces, tabs, and multiple line breaks."""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                text = run.text
                text = text.replace('\t', ' ')
                text = re.sub(r' {2,}', ' ', text)
                text = re.sub(r'\n{3,}', '\n\n', text)
                run.text = text


def _format_chapters(doc):
    """Format chapter headings with page breaks and center alignment."""
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Heading 1':
            para.paragraph_format.page_break_before = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for run in para.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(18)
                run.bold = True


def _apply_body_formatting(doc, line_spacing):
    """Apply body text formatting to all regular paragraphs."""
    for para in doc.paragraphs:
        if para.style and para.style.name in ['Normal', 'Body Text', 'Body']:
            para.paragraph_format.line_spacing = line_spacing
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Inches(0.25)
            
            for run in para.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(11)


def _insert_toc(doc):
    """Insert a Word-compatible dynamic table of contents."""
    body = doc.element.body
    
    toc_elements = []
    
    toc_title = doc.add_paragraph()
    toc_title_run = toc_title.add_run("Table of Contents")
    toc_title_run.bold = True
    toc_title_run.font.size = Pt(18)
    toc_title_run.font.name = 'Georgia'
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_elements.append(toc_title._element)
    
    spacer = doc.add_paragraph()
    toc_elements.append(spacer._element)
    
    toc_para = doc.add_paragraph()
    toc_elements.append(toc_para._element)
    
    _add_toc_field(toc_para)
    
    page_break_para = doc.add_paragraph()
    run = page_break_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    toc_elements.append(page_break_para._element)
    
    front_matter_count = 0
    for i, element in enumerate(body):
        if element.tag.endswith('}p'):
            front_matter_count += 1
            if front_matter_count > 30:
                break
    
    insert_position = min(front_matter_count, len(body) - 1)
    
    for element in reversed(toc_elements):
        if insert_position < len(body):
            body.insert(insert_position, element)
        else:
            body.append(element)


def _add_toc_field(paragraph):
    """Add TOC field codes to a paragraph for Word compatibility."""
    run = paragraph.add_run()
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)
    
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)
    
    run3 = paragraph.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar_separate)
    
    run4 = paragraph.add_run("Right-click and select 'Update Field' to generate TOC")
    run4.italic = True
    run4.font.size = Pt(10)
    
    run5 = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar_end)
