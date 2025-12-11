"""
Front Matter Service
Inserts title page, copyright page, and dedication page into the document.
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def insert_front_matter(doc, title="Untitled", author="Author Name"):
    """
    Insert front matter pages at the beginning of the document.
    
    Args:
        doc: python-docx Document object
        title: Book title
        author: Author name
    
    Returns:
        Modified Document object
    """
    body = doc.element.body
    first_element = body[0] if len(body) > 0 else None
    
    front_matter_elements = []
    
    title_elements = _create_title_page(doc, title, author)
    front_matter_elements.extend(title_elements)
    
    copyright_elements = _create_copyright_page(doc, title, author)
    front_matter_elements.extend(copyright_elements)
    
    dedication_elements = _create_dedication_page(doc)
    front_matter_elements.extend(dedication_elements)
    
    for element in reversed(front_matter_elements):
        if first_element is not None:
            first_element.addprevious(element)
        else:
            body.append(element)
    
    return doc


def _create_title_page(doc, title, author):
    """Create title page elements."""
    elements = []
    
    for _ in range(8):
        p = doc.add_paragraph()
        elements.append(p._element)
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.name = 'Georgia'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elements.append(title_para._element)
    
    spacer = doc.add_paragraph()
    elements.append(spacer._element)
    spacer2 = doc.add_paragraph()
    elements.append(spacer2._element)
    
    author_para = doc.add_paragraph()
    author_run = author_para.add_run(author)
    author_run.font.size = Pt(18)
    author_run.font.name = 'Georgia'
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elements.append(author_para._element)
    
    page_break = doc.add_paragraph()
    run = page_break.add_run()
    run._element.append(_create_page_break())
    elements.append(page_break._element)
    
    return elements


def _create_copyright_page(doc, title, author):
    """Create copyright page elements."""
    elements = []
    current_year = datetime.now().year
    
    for _ in range(6):
        p = doc.add_paragraph()
        elements.append(p._element)
    
    copyright_text = f"""Copyright © {current_year} {author}

All rights reserved.

No part of this publication may be reproduced, distributed, or transmitted in any form or by any means, including photocopying, recording, or other electronic or mechanical methods, without the prior written permission of the publisher, except in the case of brief quotations embodied in critical reviews and certain other noncommercial uses permitted by copyright law.

{title}

First Edition

Printed in the United States of America"""

    for line in copyright_text.split('\n'):
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(10)
        run.font.name = 'Georgia'
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elements.append(para._element)
    
    page_break = doc.add_paragraph()
    run = page_break.add_run()
    run._element.append(_create_page_break())
    elements.append(page_break._element)
    
    return elements


def _create_dedication_page(doc):
    """Create dedication page elements."""
    elements = []
    
    for _ in range(10):
        p = doc.add_paragraph()
        elements.append(p._element)
    
    dedication_para = doc.add_paragraph()
    dedication_run = dedication_para.add_run("For someone special…")
    dedication_run.italic = True
    dedication_run.font.size = Pt(14)
    dedication_run.font.name = 'Georgia'
    dedication_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elements.append(dedication_para._element)
    
    page_break = doc.add_paragraph()
    run = page_break.add_run()
    run._element.append(_create_page_break())
    elements.append(page_break._element)
    
    return elements


def _create_page_break():
    """Create a page break XML element."""
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br
